from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Table, TableStyle


@dataclass
class Segment:
    index: int
    original: str


def parse_docx_segments(content: bytes) -> tuple[DocumentType, list[Segment]]:
    doc = Document(BytesIO(content))
    segments: list[Segment] = []
    i = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            segments.append(Segment(index=i, original=text))
            i += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if text:
                    segments.append(Segment(index=i, original=f"[TABLE] {text}"))
                    i += 1

    return doc, segments


def chunk_segments(segments: list[Segment], max_chars: int) -> list[list[Segment]]:
    chunks: list[list[Segment]] = []
    current: list[Segment] = []
    current_len = 0

    for seg in segments:
        seg_len = len(seg.original) + 20
        if current and current_len + seg_len > max_chars:
            chunks.append(current)
            current = []
            current_len = 0
        current.append(seg)
        current_len += seg_len

    if current:
        chunks.append(current)
    return chunks


def build_bilingual_docx(rows: Iterable[tuple[str, str]], output_path: Path) -> None:
    doc = Document()
    doc.add_heading("Bilingual Legal Translation", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Original"
    hdr[1].text = "English"

    for original, translated in rows:
        row = table.add_row().cells
        row[0].text = original
        row[1].text = translated

    doc.save(str(output_path))


def build_bilingual_pdf(rows: Iterable[tuple[str, str]], output_path: Path) -> None:
    styles = getSampleStyleSheet()
    cell_style = ParagraphStyle(
        "Cell",
        parent=styles["Normal"],
        fontSize=9,
        leading=12,
        spaceAfter=4,
    )

    data = [["Original", "English"]]
    for original, translated in rows:
        data.append([Paragraph(original, cell_style), Paragraph(translated, cell_style)])

    pdf = SimpleDocTemplate(str(output_path), pagesize=A4, leftMargin=1.2 * cm, rightMargin=1.2 * cm)
    table = Table(data, colWidths=[9 * cm, 9 * cm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    pdf.build([table])
